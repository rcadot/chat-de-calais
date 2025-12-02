# tests/conftest.py (AJOUT)
"""Fixtures partagées pour tous les tests."""
import pytest
import tempfile
import os
import shutil
from pathlib import Path
from unittest.mock import Mock, MagicMock
from datetime import datetime

@pytest.fixture
def temp_dir():
    """Crée un répertoire temporaire."""
    temp = tempfile.mkdtemp()
    yield temp
    shutil.rmtree(temp)

@pytest.fixture
def temp_documents_dir(temp_dir):
    """Crée un répertoire de documents temporaire avec des fichiers de test."""
    docs_dir = os.path.join(temp_dir, "documents")
    os.makedirs(docs_dir)
    
    # ✅ TXT
    with open(os.path.join(docs_dir, "test.txt"), "w", encoding="utf-8") as f:
        f.write("Ceci est un document de test sur le ZAN.\nDeuxième ligne.")
    
    # ✅ MD
    with open(os.path.join(docs_dir, "test.md"), "w", encoding="utf-8") as f:
        f.write("# Titre\n\nContenu markdown sur Mad et Moselle.")
    
    # ✅ DOCX (si python-docx disponible)
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument()
        doc.add_heading('Document Test', 0)
        doc.add_paragraph('Contenu DOCX de test sur le ZAN.')
        doc.save(os.path.join(docs_dir, "test.docx"))
    except ImportError:
        pass
    
    # ✅ PDF (si reportlab disponible)
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        pdf_path = os.path.join(docs_dir, "test.pdf")
        c = canvas.Canvas(pdf_path, pagesize=letter)
        c.drawString(100, 750, "Document PDF test")
        c.drawString(100, 730, "Contenu sur le ZAN")
        c.save()
    except ImportError:
        pass
    
    # ✅ ODT (si odfpy disponible)
    try:
        from odf.opendocument import OpenDocumentText
        from odf.text import P
        textdoc = OpenDocumentText()
        p = P(text="Document ODT de test.")
        textdoc.text.addElement(p)
        p2 = P(text="Contenu sur le ZAN.")
        textdoc.text.addElement(p2)
        textdoc.save(os.path.join(docs_dir, "test.odt"))
    except ImportError:
        pass
    
    # ✅ HTML
    html_content = """
    <!DOCTYPE html>
    <html>
    <head><title>Test</title></head>
    <body>
        <h1>Test HTML</h1>
        <p>Contenu sur le ZAN.</p>
    </body>
    </html>
    """
    with open(os.path.join(docs_dir, "test.html"), "w", encoding="utf-8") as f:
        f.write(html_content)
    
    return docs_dir

@pytest.fixture
def mock_embeddings():
    """Mock des embeddings Albert."""
    mock = MagicMock()
    mock.embed_documents.return_value = [[0.1, 0.2, 0.3]] * 5
    mock.embed_query.return_value = [0.1, 0.2, 0.3]
    return mock

@pytest.fixture
def mock_llm():
    """Mock du LLM Albert."""
    mock = MagicMock()
    mock_response = MagicMock()
    mock_response.content = "Ceci est une réponse générée par le LLM."
    mock.invoke.return_value = mock_response
    
    # Pour streaming
    mock_chunks = [
        MagicMock(content="Ceci "),
        MagicMock(content="est "),
        MagicMock(content="une "),
        MagicMock(content="réponse.")
    ]
    mock.stream.return_value = iter(mock_chunks)
    
    return mock

@pytest.fixture
def sample_documents():
    """Documents LangChain de test."""
    from langchain_core.documents import Document
    return [
        Document(page_content="Contenu 1 sur le ZAN", metadata={"source": "doc1.pdf"}),
        Document(page_content="Contenu 2 sur Mad et Moselle", metadata={"source": "doc2.odt"}),
        Document(page_content="Contenu 3 général", metadata={"source": "doc3.txt"}),
    ]

@pytest.fixture
def mock_retriever(sample_documents):
    """Mock du retriever."""
    mock = MagicMock()
    mock.invoke.return_value = sample_documents
    return mock

@pytest.fixture
def temp_db_path(temp_dir):
    """Chemin vers une base ChromaDB temporaire."""
    return os.path.join(temp_dir, "test_chroma_db")

@pytest.fixture
def temp_logs_db(temp_dir):
    """Chemin vers une base de logs temporaire."""
    return os.path.join(temp_dir, "test_logs.db")

@pytest.fixture(autouse=True)
def setup_test_env(monkeypatch, temp_dir):
    """Configure l'environnement pour les tests."""
    monkeypatch.setenv("ALBERT_API_KEY", "test_key_12345")
    
    # Importer config après avoir set les env vars
    import config
    monkeypatch.setattr(config, "VERBOSE", False)  # Désactiver logs pendant tests
